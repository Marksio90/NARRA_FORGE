"""
Export service for book formats
Exports completed books to DOCX, EPUB, PDF, and Markdown
"""

from typing import List
from sqlalchemy.orm import Session
import logging
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from ebooklib import epub

from app.config import settings
from app.models.project import Project
from app.models.chapter import Chapter

logger = logging.getLogger(__name__)


async def export_to_format(db: Session, project_id: int, format: str) -> str:
    """
    Export project to specified format
    
    Args:
        db: Database session
        project_id: Project ID
        format: Output format (docx, epub, pdf, markdown)
    
    Returns:
        Path to exported file
    """
    # Get project and chapters
    project = db.query(Project).filter(Project.id == project_id).first()
    chapters = db.query(Chapter).filter(
        Chapter.project_id == project_id
    ).order_by(Chapter.number).all()
    
    if not project or not chapters:
        raise ValueError("Project or chapters not found")
    
    # Create output directory
    output_dir = Path(settings.OUTPUT_DIR) / str(project_id)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate filename
    safe_name = "".join(c for c in project.name if c.isalnum() or c in (' ', '-', '_')).strip()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Export based on format
    if format == "docx":
        file_path = output_dir / f"{safe_name}_{timestamp}.docx"
        _export_docx(project, chapters, file_path)
    elif format == "epub":
        file_path = output_dir / f"{safe_name}_{timestamp}.epub"
        _export_epub(project, chapters, file_path)
    elif format == "pdf":
        file_path = output_dir / f"{safe_name}_{timestamp}.pdf"
        _export_pdf(project, chapters, file_path)
    elif format == "markdown":
        file_path = output_dir / f"{safe_name}_{timestamp}.md"
        _export_markdown(project, chapters, file_path)
    else:
        raise ValueError(f"Unsupported format: {format}")
    
    logger.info(f"Exported project {project_id} to {format}: {file_path}")
    
    return str(file_path)


def _export_docx(project: Project, chapters: List[Chapter], file_path: Path):
    """Export to Microsoft Word DOCX with professional book formatting"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Set up styles for professional book formatting
    styles = doc.styles

    # Body text style (for narrative paragraphs)
    try:
        body_style = styles['Body Text']
    except KeyError:
        body_style = styles.add_style('Body Text', 1)  # 1 = paragraph style

    # Configure body text for book-like appearance
    body_font = body_style.font
    body_font.name = 'Times New Roman'  # Classic book font
    body_font.size = Pt(12)

    # First line indent (classic book style)
    body_para = body_style.paragraph_format
    body_para.first_line_indent = Inches(0.3)  # Half-inch indent
    body_para.line_spacing = 1.15  # Comfortable reading
    body_para.space_after = Pt(0)  # No space between paragraphs (book style)
    body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justified text

    # Dialogue style (for paragraphs starting with em dash)
    try:
        dialogue_style = styles['Dialogue']
    except KeyError:
        dialogue_style = styles.add_style('Dialogue', 1)

    dialogue_font = dialogue_style.font
    dialogue_font.name = 'Times New Roman'
    dialogue_font.size = Pt(12)

    dialogue_para = dialogue_style.paragraph_format
    dialogue_para.first_line_indent = Inches(0.3)
    dialogue_para.line_spacing = 1.15
    dialogue_para.space_after = Pt(0)
    dialogue_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Chapter title style
    try:
        chapter_style = styles['Heading 1']
        chapter_font = chapter_style.font
        chapter_font.name = 'Times New Roman'
        chapter_font.size = Pt(18)
        chapter_font.bold = True
        chapter_para = chapter_style.paragraph_format
        chapter_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chapter_para.space_before = Pt(72)  # Space before chapter title
        chapter_para.space_after = Pt(24)
    except KeyError:
        logger.warning("Style 'Heading 1' not found in document, chapter titles will use default style")
    except Exception as e:
        logger.warning(f"Failed to apply chapter title style: {e}")

    # Add page numbers to footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page number field
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

    # Title page
    title = doc.add_heading(project.name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_font = title.runs[0].font
    title_font.name = 'Times New Roman'
    title_font.size = Pt(24)
    title_font.bold = True

    doc.add_paragraph()

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(f"Gatunek: {project.genre.value.title()}\n")
    info_run.font.name = 'Times New Roman'
    info_run.font.size = Pt(11)

    date_run = info_para.add_run(f"Wygenerowano: {datetime.now().strftime('%d.%m.%Y')}\n\n")
    date_run.font.name = 'Times New Roman'
    date_run.font.size = Pt(11)

    credit_run = info_para.add_run("Stworzone przez NarraForge")
    credit_run.font.name = 'Times New Roman'
    credit_run.font.size = Pt(10)
    credit_run.italic = True

    doc.add_page_break()

    # Chapters
    for chapter in chapters:
        if chapter.content:
            # Chapter title
            chapter_heading = doc.add_heading(f"Rozdział {chapter.number}", level=1)
            if chapter.title:
                title_para = doc.add_paragraph(chapter.title)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_para.runs[0]
                title_run.font.name = 'Times New Roman'
                title_run.font.size = Pt(14)
                title_run.italic = True
                title_para.space_after = Pt(24)

            # Chapter content with smart formatting
            paragraphs = chapter.content.split('\n\n')
            for i, paragraph_text in enumerate(paragraphs):
                paragraph_text = paragraph_text.strip()
                if not paragraph_text:
                    continue

                # Detect if paragraph is dialogue (starts with em dash)
                is_dialogue = paragraph_text.startswith('—') or paragraph_text.startswith('-')

                # Choose appropriate style
                if is_dialogue:
                    para = doc.add_paragraph(paragraph_text, style='Dialogue')
                else:
                    para = doc.add_paragraph(paragraph_text, style='Body Text')

                # First paragraph of chapter shouldn't have indent (book convention)
                if i == 0:
                    para.paragraph_format.first_line_indent = Pt(0)

            doc.add_page_break()

    # Save
    doc.save(file_path)
    logger.info(f"DOCX exported with professional formatting: {file_path}")


def _export_epub(project: Project, chapters: List[Chapter], file_path: Path):
    """Export to EPUB e-book format with professional book styling"""
    book = epub.EpubBook()

    # Metadata
    book.set_identifier(f"narraforge_{project.id}")
    book.set_title(project.name)
    book.set_language('pl')
    book.add_author('NarraForge AI')

    # Add CSS for professional book styling
    book_css = '''
    @namespace epub "http://www.idpf.org/2007/ops";

    body {
        font-family: "Times New Roman", Georgia, serif;
        font-size: 1.1em;
        line-height: 1.5;
        text-align: justify;
        margin: 1em;
    }

    h1 {
        font-family: "Times New Roman", Georgia, serif;
        font-size: 1.8em;
        text-align: center;
        margin-top: 3em;
        margin-bottom: 1em;
        font-weight: bold;
    }

    h2 {
        font-family: "Times New Roman", Georgia, serif;
        font-size: 1.3em;
        text-align: center;
        margin-bottom: 2em;
        font-style: italic;
        font-weight: normal;
    }

    p {
        margin: 0;
        text-indent: 1.5em;
        orphans: 2;
        widows: 2;
    }

    p.first {
        text-indent: 0;  /* First paragraph of chapter has no indent */
    }

    p.dialogue {
        text-indent: 1.5em;
    }

    .chapter-break {
        page-break-before: always;
    }
    '''

    nav_css = epub.EpubItem(
        uid="style_nav",
        file_name="style/nav.css",
        media_type="text/css",
        content=book_css
    )
    book.add_item(nav_css)

    # Add chapters
    epub_chapters = []
    for chapter in chapters:
        if chapter.content:
            c = epub.EpubHtml(
                title=f"Rozdział {chapter.number}",
                file_name=f"chap_{chapter.number:03d}.xhtml",
                lang='pl'
            )

            content = f'<div class="chapter-break"><h1>Rozdział {chapter.number}</h1>'
            if chapter.title:
                content += f"<h2>{chapter.title}</h2>"

            # Process paragraphs with proper styling
            paragraphs = [p.strip() for p in chapter.content.split('\n\n') if p.strip()]
            for i, paragraph_text in enumerate(paragraphs):
                # Detect dialogue
                is_dialogue = paragraph_text.startswith('—') or paragraph_text.startswith('-')

                # Apply appropriate class
                if i == 0:
                    # First paragraph - no indent
                    content += f'<p class="first">{paragraph_text}</p>'
                elif is_dialogue:
                    content += f'<p class="dialogue">{paragraph_text}</p>'
                else:
                    content += f'<p>{paragraph_text}</p>'

            content += '</div>'
            c.content = content
            c.add_item(nav_css)  # Link CSS to chapter
            book.add_item(c)
            epub_chapters.append(c)
    
    # Define TOC
    book.toc = tuple(epub_chapters)
    
    # Add navigation files
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    
    # Spine
    book.spine = ['nav'] + epub_chapters
    
    # Write
    epub.write_epub(file_path, book)
    logger.info(f"EPUB exported: {file_path}")


def _export_pdf(project: Project, chapters: List[Chapter], file_path: Path):
    """Export to PDF with Polish character support"""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    doc = SimpleDocTemplate(str(file_path), pagesize=letter)
    styles = getSampleStyleSheet()

    # Register Unicode-capable font (Helvetica supports Polish characters)
    # Note: Helvetica is built-in and supports extended Latin characters
    font_name = 'Helvetica'

    # Custom styles with Unicode-capable font
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        alignment=TA_CENTER,
        spaceAfter=30,
        fontName=font_name
    )

    chapter_style = ParagraphStyle(
        'ChapterTitle',
        parent=styles['Heading2'],
        fontSize=18,
        spaceAfter=12,
        fontName=font_name
    )

    body_style = ParagraphStyle(
        'Body',
        parent=styles['BodyText'],
        alignment=TA_JUSTIFY,
        fontSize=12,
        leading=16,
        fontName=font_name,  # Helvetica supports Polish characters
        firstLineIndent=18,  # First line indent (book style)
        leftIndent=0,
        rightIndent=0,
        spaceAfter=0,  # No space between paragraphs
        spaceBefore=0
    )

    # Dialogue style (for lines starting with em dash)
    dialogue_style = ParagraphStyle(
        'Dialogue',
        parent=styles['BodyText'],
        alignment=TA_JUSTIFY,
        fontSize=12,
        leading=16,
        fontName=font_name,
        firstLineIndent=18,
        leftIndent=0,
        rightIndent=0,
        spaceAfter=0,
        spaceBefore=0
    )

    # First paragraph style (no indent, book convention)
    first_para_style = ParagraphStyle(
        'FirstParagraph',
        parent=styles['BodyText'],
        alignment=TA_JUSTIFY,
        fontSize=12,
        leading=16,
        fontName=font_name,
        firstLineIndent=0,  # No indent for first paragraph
        leftIndent=0,
        rightIndent=0,
        spaceAfter=0,
        spaceBefore=0
    )
    
    # Build content
    content = []
    
    # Title page
    content.append(Paragraph(project.name, title_style))
    content.append(Spacer(1, 0.2*Inches))
    content.append(Paragraph(f"Gatunek: {project.genre.value.title()}", styles['Normal']))
    content.append(Spacer(1, 0.5*Inches))
    content.append(Paragraph("Stworzone przez NarraForge", styles['Normal']))
    content.append(PageBreak())
    
    # Chapters
    for chapter in chapters:
        if chapter.content:
            content.append(Paragraph(f"Rozdział {chapter.number}", chapter_style))
            if chapter.title:
                title_style_italic = ParagraphStyle(
                    'ChapterTitleItalic',
                    parent=styles['Heading3'],
                    fontName='Helvetica-Oblique',  # Helvetica italic for Polish support
                    alignment=TA_CENTER,
                    fontSize=14
                )
                content.append(Paragraph(chapter.title, title_style_italic))

            content.append(Spacer(1, 0.3*Inches))

            # Chapter content with professional formatting
            paragraphs = [p.strip() for p in chapter.content.split('\n\n') if p.strip()]
            for i, paragraph_text in enumerate(paragraphs):
                # Detect if dialogue (starts with em dash)
                is_dialogue = paragraph_text.startswith('—') or paragraph_text.startswith('-')

                # Choose style: first paragraph has no indent
                if i == 0:
                    current_style = first_para_style
                elif is_dialogue:
                    current_style = dialogue_style
                else:
                    current_style = body_style

                try:
                    # Ensure proper encoding for Polish characters
                    safe_text = paragraph_text.encode('utf-8', 'replace').decode('utf-8')
                    content.append(Paragraph(safe_text, current_style))
                    content.append(Spacer(1, 0.05*Inches))  # Small space between paragraphs
                except Exception as e:
                    logger.warning(f"Failed to add paragraph to PDF: {e}")
                    # Add as plain text if Paragraph fails
                    content.append(Paragraph(paragraph_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), current_style))
                    content.append(Spacer(1, 0.05*Inches))

            content.append(PageBreak())
    
    # Build PDF
    doc.build(content)
    logger.info(f"PDF exported: {file_path}")


def _export_markdown(project: Project, chapters: List[Chapter], file_path: Path):
    """Export to Markdown with Polish character support"""
    try:
        with open(file_path, 'w', encoding='utf-8', errors='replace') as f:
            # Title
            f.write(f"# {project.name}\n\n")
            f.write(f"**Gatunek:** {project.genre.value.title()}  \n")
            f.write(f"**Wygenerowano:** {datetime.now().strftime('%d.%m.%Y')}  \n")
            f.write(f"**Platforma:** NarraForge\n\n")
            f.write("---\n\n")

            # Chapters
            for chapter in chapters:
                if chapter.content:
                    f.write(f"## Rozdział {chapter.number}\n\n")
                    if chapter.title:
                        f.write(f"### {chapter.title}\n\n")

                    # Ensure content is properly formatted
                    content = chapter.content
                    if not content.endswith('\n'):
                        content += '\n'

                    f.write(content)
                    f.write("\n---\n\n")

        logger.info(f"Markdown exported: {file_path}")
    except Exception as e:
        logger.error(f"Markdown export failed: {e}", exc_info=True)
        raise
